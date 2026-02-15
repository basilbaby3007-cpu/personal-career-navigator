from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn
from datetime import datetime
import random
import re
import io

# Create the app
app = FastAPI()

# Allow frontend to connect
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Define what data we expect
class UserProfile(BaseModel):
    resume_text: str = ""
    dream_role: str = "ml_engineer"
    hours_per_week: int = 15
    current_skills: list = []

# ========== JOB REQUIREMENTS DATABASE ==========
job_requirements = {
    "ml_engineer": {
        "required": ["python", "tensorflow", "pytorch", "sql", "machine learning", "data structures", "algorithms"],
        "preferred": ["docker", "aws", "kubernetes", "spark", "computer vision", "nlp"],
        "projects": ["Build an image classifier", "Create a recommendation system", "Fine-tune a language model"]
    },
    "fullstack_developer": {
        "required": ["javascript", "react", "node.js", "html", "css", "git", "sql", "rest api"],
        "preferred": ["typescript", "docker", "aws", "graphql", "mongodb", "next.js"],
        "projects": ["Build a full-stack CRUD app", "Create a real-time chat app", "E-commerce website"]
    },
    "data_scientist": {
        "required": ["python", "sql", "pandas", "numpy", "machine learning", "statistics", "data visualization"],
        "preferred": ["spark", "tensorflow", "aws", "tableau", "big data", "scikit-learn"],
        "projects": ["Exploratory data analysis", "Predictive modeling", "Data visualization dashboard"]
    },
    "devops_engineer": {
        "required": ["linux", "docker", "kubernetes", "aws", "jenkins", "git", "python", "ci/cd"],
        "preferred": ["terraform", "ansible", "prometheus", "grafana", "bash"],
        "projects": ["CI/CD pipeline setup", "Infrastructure as Code", "Monitoring system"]
    },
    "cybersecurity_analyst": {
        "required": ["networking", "linux", "python", "security", "firewalls", "encryption"],
        "preferred": ["ethical hacking", "cloud security", "penetration testing", "siem"],
        "projects": ["Security audit tool", "Network scanner", "Incident response plan"]
    },
    "product_manager": {
        "required": ["agile", "user research", "analytics", "communication", "roadmapping", "product strategy"],
        "preferred": ["technical background", "figma", "sql", "a/b testing", "jira"],
        "projects": ["Product requirements doc", "User research study", "Product roadmap"]
    }
}

# ========== RESOURCES DATABASE ==========
learning_resources = {
    "python": "https://www.python.org/about/gettingstarted/",
    "javascript": "https://javascript.info/",
    "react": "https://react.dev/learn",
    "node.js": "https://nodejs.org/en/learn",
    "tensorflow": "https://www.tensorflow.org/learn",
    "pytorch": "https://pytorch.org/tutorials/",
    "docker": "https://docs.docker.com/get-started/",
    "kubernetes": "https://kubernetes.io/docs/tutorials/",
    "aws": "https://aws.amazon.com/training/",
    "git": "https://git-scm.com/doc",
    "sql": "https://www.w3schools.com/sql/",
    "machine learning": "https://www.coursera.org/learn/machine-learning",
    "linux": "https://linuxjourney.com/",
    "html": "https://www.w3schools.com/html/",
    "css": "https://www.w3schools.com/css/",
    "mongodb": "https://university.mongodb.com/",
    "pandas": "https://pandas.pydata.org/docs/",
    "numpy": "https://numpy.org/doc/stable/",
    "algorithms": "https://www.coursera.org/learn/algorithms",
    "rest api": "https://restfulapi.net/",
    "data visualization": "https://www.tableau.com/learn",
}

# ========== SKILL DATABASE ==========
# Create a list of all possible skills from requirements
all_skills = []
for role, data in job_requirements.items():
    all_skills.extend(data["required"])
    all_skills.extend(data["preferred"])
all_skills = list(set(all_skills))  # Remove duplicates

# Add learning resources skills
all_skills.extend(list(learning_resources.keys()))
all_skills = list(set(all_skills))

# Create skill synonyms for better matching
skill_synonyms = {
    "python": ["python", "py", "python3", "python 3"],
    "javascript": ["javascript", "js", "ecmascript", "typescript", "ts"],
    "react": ["react", "reactjs", "react.js", "jsx"],
    "node.js": ["node", "nodejs", "node.js", "express"],
    "machine learning": ["machine learning", "ml", "ai", "artificial intelligence", "deep learning", "neural networks"],
    "sql": ["sql", "mysql", "postgresql", "sqlite", "database", "relational", "t-sql"],
    "docker": ["docker", "container", "containerization", "dockerfile"],
    "aws": ["aws", "amazon web services", "amazon aws", "ec2", "s3", "lambda", "cloud"],
    "git": ["git", "github", "gitlab", "bitbucket", "version control", "vcs"],
    "html": ["html", "html5"],
    "css": ["css", "css3", "styling", "sass", "scss", "less"],
    "data structures": ["data structures", "dsa", "arrays", "linked lists", "trees", "graphs", "hashmap"],
    "algorithms": ["algorithms", "algorithm", "sorting", "searching", "dynamic programming"],
    "rest api": ["rest", "api", "rest api", "restful", "http"],
    "linux": ["linux", "unix", "ubuntu", "bash", "shell", "command line"],
    "pandas": ["pandas", "data analysis", "dataframe"],
    "numpy": ["numpy", "numerical python", "scientific computing"],
    "ci/cd": ["ci/cd", "cicd", "continuous integration", "continuous deployment", "jenkins", "gitlab ci"],
    "tensorflow": ["tensorflow", "tf", "neural networks"],
    "pytorch": ["pytorch", "pth"],
    "kubernetes": ["kubernetes", "k8s", "orchestration"],
    "mongodb": ["mongodb", "nosql", "document database"],
    "react": ["react", "reactjs", "jsx"],
    "fullstack": ["fullstack", "full stack", "full-stack", "frontend", "backend"],
    "agile": ["agile", "scrum", "sprint", "kanban", "project management"],
    "statistics": ["statistics", "statistical", "probability", "inference"],
    "security": ["security", "cybersecurity", "encryption", "authentication"],
    "networking": ["networking", "network", "tcp/ip"],
    "ansible": ["ansible", "infrastructure automation"],
    "terraform": ["terraform", "infrastructure as code", "iac"],
    "prometheus": ["prometheus", "monitoring"],
    "grafana": ["grafana", "dashboards"],
}


# Test route
@app.get("/")
def home():
    return {
        "message": "Career Navigator API is LIVE! 🚀",
        "version": "2.0",
        "features": ["Skill Analysis", "Gap Detection", "Smart Roadmap", "Progress Tracking"]
    }

# ========== FILE UPLOAD & EXTRACTION =========
@app.post("/api/extract-resume")
async def extract_resume(file: UploadFile = File(...)):
    """Extract text from uploaded resume (TXT, PDF, DOCX)"""
    try:
        filename = file.filename.lower()
        content = await file.read()
        
        extracted_text = ""
        
        # Handle TXT files
        if filename.endswith('.txt'):
            extracted_text = content.decode('utf-8', errors='ignore')
        
        # Handle PDF files with fallback
        elif filename.endswith('.pdf'):
            pdf_success = False
            
            # Try pdfplumber first (more reliable)
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            extracted_text += text + "\n"
                pdf_success = True
            except Exception as e1:
                # Fallback to PyPDF2
                try:
                    import PyPDF2
                    pdf_file = io.BytesIO(content)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            extracted_text += text + "\n"
                    pdf_success = True
                except Exception as e2:
                    return {
                        "success": False,
                        "error": f"PDF extraction failed. Tried pdfplumber and PyPDF2. Errors: {str(e1)[:100]}",
                        "extracted_text": ""
                    }
        
        # Handle DOCX files
        elif filename.endswith('.docx') or filename.endswith('.doc'):
            try:
                from docx import Document
                doc = Document(io.BytesIO(content))
                extracted_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            except Exception as e:
                return {
                    "success": False,
                    "error": f"DOCX parsing failed: {str(e)[:100]}. Try converting to PDF or plain text.",
                    "extracted_text": ""
                }
        
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {filename}. Supported: .txt, .pdf, .docx",
                "extracted_text": ""
            }
        
        # Validate extracted text
        if not extracted_text or not extracted_text.strip():
            return {
                "success": False,
                "error": "No text could be extracted from the file. File may be corrupted or encrypted.",
                "extracted_text": ""
            }
        
        return {
            "success": True,
            "extracted_text": extracted_text.strip(),
            "filename": filename
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": f"Unexpected error: {str(e)[:100]}",
            "extracted_text": ""
        }

# Alternative: Handle base64 encoded files to avoid multipart requirement
class FileUploadBase64(BaseModel):
    filename: str
    content_base64: str

@app.post("/api/extract-resume-base64")
async def extract_resume_base64(data: FileUploadBase64):
    """Extract text from base64 encoded file (avoids multipart dependency)"""
    try:
        import base64
        filename = data.filename.lower()
        content = base64.b64decode(data.content_base64)
        
        extracted_text = ""
        
        # Handle TXT files
        if filename.endswith('.txt'):
            extracted_text = content.decode('utf-8', errors='ignore')
        
        # Handle PDF files
        elif filename.endswith('.pdf'):
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            extracted_text += text + "\n"
            except Exception as e1:
                try:
                    import PyPDF2
                    pdf_file = io.BytesIO(content)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            extracted_text += text + "\n"
                except Exception as e2:
                    return {
                        "success": False,
                        "error": f"PDF extraction failed: {str(e1)[:80]}",
                        "extracted_text": ""
                    }
        
        # Handle DOCX files
        elif filename.endswith('.docx') or filename.endswith('.doc'):
            try:
                from docx import Document
                doc = Document(io.BytesIO(content))
                extracted_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            except Exception as e:
                return {
                    "success": False,
                    "error": f"DOCX parsing failed: {str(e)[:80]}",
                    "extracted_text": ""
                }
        
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {filename}",
                "extracted_text": ""
            }
        
        if not extracted_text or not extracted_text.strip():
            return {
                "success": False,
                "error": "No text could be extracted from the file.",
                "extracted_text": ""
            }
        
        return {
            "success": True,
            "extracted_text": extracted_text.strip(),
            "filename": filename
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": f"Error: {str(e)[:100]}",
            "extracted_text": ""
        }
# ========== IMPROVED SKILL EXTRACTION ==========
def extract_skills_from_text(text):
    """Extract skills from text using keyword matching, synonyms, and context analysis"""
    text = text.lower()
    found_skills = set()
    
    # Look for experience indicators with skills
    experience_patterns = [
        r'(\d+)\s+(?:years?|yrs?)?\s+(?:of\s+)?(?:experience\s+)?(?:with\s+)?([a-zA-Z\s\+\#\.\-]+)',
        r'expertise\s+in\s+([a-zA-Z\s\+\#\.\-]+)',
        r'proficient\s+(?:in|with)\s+([a-zA-Z\s\+\#\.\-]+)',
        r'skilled\s+(?:in|with)\s+([a-zA-Z\s\+\#\.\-]+)',
        r'(?:built|developed|created|designed|implemented)\s+(?:[a-zA-Z\s]+?\s+)?(?:with|using|in)\s+([a-zA-Z\s\+\#\.\-]+)',
    ]
    
    # Direct skill matching
    for skill in all_skills:
        skill_lower = skill.lower()
        
        # Direct exact or substring match
        if skill_lower in text:
            found_skills.add(skill)
            continue
        
        # Check synonyms
        if skill in skill_synonyms:
            for synonym in skill_synonyms[skill]:
                if synonym.lower() in text:
                    found_skills.add(skill)
                    break
    
    # Context-aware matching: look for technical terms mentioned with action words
    action_words = ['built', 'developed', 'created', 'designed', 'implemented', 'worked', 'used', 'experienced']
    tech_terms = text.split()
    
    for i, word in enumerate(tech_terms):
        # Look for action verb + tech skill pattern
        if word in action_words and i + 1 < len(tech_terms):
            next_words = ' '.join(tech_terms[i+1:min(i+5, len(tech_terms))])
            for skill in all_skills:
                if skill.lower() in next_words.lower():
                    found_skills.add(skill)
    
    return list(found_skills)


# ========== MAIN ROADMAP FUNCTION ==========
@app.post("/api/roadmap")
def create_roadmap(profile: UserProfile):
    print("\n" + "="*60)
    print("📥 NEW ROADMAP REQUEST")
    print("="*60)
    print(f"📋 Dream Role: {profile.dream_role}")
    print(f"⏰ Hours/Week: {profile.hours_per_week}")
    print(f"📄 Resume Length: {len(profile.resume_text)} characters")
    
    # Step 1: Extract skills from resume using improved method
    user_skills = extract_skills_from_text(profile.resume_text)
    
    print(f"\n🔍 Skills extracted from resume: {user_skills}")
    
    # Add skills from input field
    input_skills = []
    for skill in profile.current_skills:
        skill = skill.lower().strip()
        if skill:
            input_skills.append(skill)
            # Check if this skill matches any in our database
            matched = False
            for db_skill in all_skills:
                if skill in db_skill.lower() or db_skill.lower() in skill:
                    if db_skill not in user_skills:
                        user_skills.append(db_skill)
                    matched = True
                    break
            if not matched and skill not in user_skills:
                user_skills.append(skill)  # Add as-is if no match
    
    user_skills = list(set(user_skills))  # Remove duplicates
    print(f"📝 Skills from input field: {input_skills}")
    print(f"✅ Total unique skills found: {len(user_skills)} - {user_skills}")

    
    # Step 2: Get requirements for dream role
    role = profile.dream_role
    if role not in job_requirements:
        role = "ml_engineer"
    
    requirements = job_requirements[role]
    
    # Step 3: Find skill gaps with better matching
    missing_required = []
    missing_preferred = []
    
    # Helper function for intelligent skill matching
    def skill_matches(required_skill, user_skills_list):
        """Check if required skill is covered by user skills with fuzzy matching"""
        required_lower = required_skill.lower()
        
        for user_skill in user_skills_list:
            user_lower = user_skill.lower()
            
            # Exact match
            if required_lower == user_lower:
                return True
            
            # Substring match (e.g., 'python' in 'python3')
            if required_lower in user_lower or user_lower in required_lower:
                return True
            
            # Synonym checking
            if required_skill in skill_synonyms:
                for syn in skill_synonyms[required_skill]:
                    if syn.lower() in user_lower:
                        return True
            
            # Check if user skill has a synonym that matches required
            for skill_key, syns in skill_synonyms.items():
                if user_lower in [s.lower() for s in syns] and skill_key.lower() == required_lower:
                    return True
        
        return False
    
    for skill in requirements["required"]:
        if not skill_matches(skill, user_skills):
            missing_required.append(skill)
    
    for skill in requirements["preferred"]:
        if not skill_matches(skill, user_skills) and skill not in missing_required:
            missing_preferred.append(skill)

    
    print(f"❌ Missing required: {missing_required}")
    print(f"⚠️ Missing preferred: {missing_preferred}")
    
    # Step 4: Generate SMART roadmap based on gaps
    roadmap = {
        "weeks": [],
        "hours_per_week": profile.hours_per_week,
        "start_date": datetime.now().strftime("%Y-%m-%d"),
        "skill_gaps": {
            "required": missing_required,
            "preferred": missing_preferred[:3]
        },
        "analysis": {
            "skills_found": user_skills,
            "total_skills": len(user_skills),
            "gaps_count": len(missing_required) + len(missing_preferred)
        }
    }
    
    # If no gaps found, add some default learning
    if len(missing_required) == 0 and len(missing_preferred) == 0:
        missing_required = ["advanced concepts", "best practices"]
    
    # Week 1: Foundation (Learn missing required skills)
    week1_tasks = []
    skills_to_learn = missing_required[:3] if missing_required else ["advanced concepts"]
    
    for i, skill in enumerate(skills_to_learn):
        resource = learning_resources.get(skill, f"https://www.google.com/search?q=learn+{skill.replace(' ', '+')}")
        week1_tasks.append({
            "day": i * 2 + 1,
            "title": f"Master {skill.title()}",
            "description": f"Learn fundamentals and best practices of {skill}",
            "resources": [resource],
            "time": f"{int(profile.hours_per_week/3)} hours",
            "type": "learning"
        })
        week1_tasks.append({
            "day": i * 2 + 2,
            "title": f"Practice {skill.title()}",
            "description": f"Build hands-on projects using {skill}",
            "resources": ["https://github.com/topics/project-ideas"],
            "time": f"{int(profile.hours_per_week/3)} hours",
            "type": "practice"
        })
    
    roadmap["weeks"].append({
        "week": 1,
        "focus": f"Core {role.replace('_', ' ').title()} Skills",
        "tasks": sorted(week1_tasks, key=lambda x: x["day"])
    })
    
    # Week 2: Project Building
    project_idea = random.choice(requirements["projects"])
    roadmap["weeks"].append({
        "week": 2,
        "focus": "Build Portfolio Project",
        "tasks": [
            {
                "day": 8,
                "title": "Project Planning",
                "description": f"Plan your project: {project_idea}",
                "resources": ["https://www.atlassian.com/agile/project-management"],
                "time": "2 hours"
            },
            {
                "day": 9,
                "title": "Setup Development Environment",
                "description": "Install tools, create repo, setup project structure",
                "resources": ["https://git-scm.com/book/en/v2/Git-Basics-Getting-a-Git-Repository"],
                "time": "1.5 hours"
            },
            {
                "day": 11,
                "title": "Core Features Implementation",
                "description": "Build main functionality of your project",
                "resources": ["https://stackoverflow.com/"],
                "time": "3 hours"
            },
            {
                "day": 14,
                "title": "MVP Complete",
                "description": "Working version with core features done",
                "resources": [],
                "time": "2.5 hours"
            }
        ]
    })
    
    # Week 3: Advanced Concepts
    week3_tasks = []
    if missing_preferred:
        for i, skill in enumerate(missing_preferred[:2]):
            resource = learning_resources.get(skill, f"https://www.google.com/search?q=learn+{skill.replace(' ', '+')}")
            week3_tasks.append({
                "day": 15 + i,
                "title": f"Advanced {skill.title()}",
                "description": f"Deep dive into {skill} with real-world applications",
                "resources": [resource],
                "time": "2.5 hours"
            })
    
    week3_tasks.extend([
        {
            "day": 18,
            "title": "Code Review & Optimization",
            "description": "Refactor code, improve performance, fix bugs",
            "resources": ["https://refactoring.guru/"],
            "time": "2 hours"
        },
        {
            "day": 20,
            "title": "Add Advanced Features",
            "description": "Enhance project with additional capabilities",
            "resources": [],
            "time": "3 hours"
        }
    ])
    
    roadmap["weeks"].append({
        "week": 3,
        "focus": "Advanced Topics & Optimization",
        "tasks": sorted(week3_tasks, key=lambda x: x["day"])
    })
    
    # Week 4: Polish & Launch
    roadmap["weeks"].append({
        "week": 4,
        "focus": "Portfolio Polish & Career Launch",
        "tasks": [
            {
                "day": 22,
                "title": "Documentation & README",
                "description": "Write comprehensive docs, setup instructions, and comments",
                "resources": ["https://www.makeareadme.com/"],
                "time": "2 hours"
            },
            {
                "day": 24,
                "title": "Create Project Demo",
                "description": "Record video demo or write case study",
                "resources": ["https://www.youtube.com/"],
                "time": "2 hours"
            },
            {
                "day": 26,
                "title": "Update LinkedIn & GitHub",
                "description": "Showcase project, update skills section, pin repositories",
                "resources": ["https://www.linkedin.com"],
                "time": "1.5 hours"
            },
            {
                "day": 28,
                "title": "Job Applications",
                "description": f"Start applying to {role.replace('_', ' ').title()} positions",
                "resources": ["https://www.linkedin.com/jobs", "https://github.com/jobs"],
                "time": "2 hours"
            },
            {
                "day": 30,
                "title": "Review & Next Steps",
                "description": "Review progress, identify next learning goals",
                "resources": [],
                "time": "1 hour"
            }
        ]
    })
    
    print("✅ Roadmap generated successfully!")
    print("="*50 + "\n")
    return {"success": True, "roadmap": roadmap}

# ========== PROGRESS ADAPTATION ==========
class AdaptRequest(BaseModel):
    completed_days: list = []
    current_roadmap: dict = {}

@app.post("/api/adapt")
def adapt_plan(request: AdaptRequest):
    """Adapt roadmap based on user progress"""
    completed_days = request.completed_days
    completion_rate = len(completed_days) / 30
    
    if completion_rate > 0.3:
        return {
            "action": "accelerate",
            "message": "You're making fast progress! Here's an accelerated path:",
            "new_tasks": [
                "Skip basic tutorials - move to advanced projects",
                "Add extra features to your project",
                "Start networking with industry professionals"
            ]
        }
    elif completion_rate < 0.1 and len(completed_days) > 0:
        return {
            "action": "simplify",
            "message": "Let's make this easier to match your pace:",
            "new_tasks": [
                "Focus on one skill at a time",
                "More foundational resources added",
                "Break down tasks into smaller steps"
            ]
        }
    else:
        return {
            "action": "maintain",
            "message": "You're on track! Keep up the good work!",
            "new_tasks": []
        }

# Run the server
if __name__ == "__main__":
    print("\n" + "="*60)
    print("✅ CAREER NAVIGATOR API V2.0 IS STARTING...")
    print("📝 Open: http://localhost:8001")
    print("🔌 Press CTRL+C to stop")
    print("="*60 + "\n")
    uvicorn.run(app, host="127.0.0.1", port=8001)